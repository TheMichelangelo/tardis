from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mikepastula/MyProjects/stem_laboratory")
OUTPUT = ROOT / "semestrova-kontrolna-stem-5-klas.docx"


def set_page_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title.paragraph_format.space_after = Pt(10)

    for name, size, before, after in [
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "Small Note" not in styles:
        note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Small Note"]
    note.base_style = normal
    note.font.name = "Arial"
    note.font.size = Pt(10)
    note.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.15


def add_paragraph(document, text="", style="Normal", bold_prefix=None, alignment=None):
    p = document.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_numbered_lines(document, items):
    for idx, item in enumerate(items, start=1):
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.1)
        lead = p.add_run(f"{idx}. ")
        lead.bold = True
        p.add_run(item)


def add_answer_lines(document, count=5):
    for _ in range(count):
        add_paragraph(document, "_" * 70)


def add_simple_table(document, headers, rows, widths=None):
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)

    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "F1F3F4")

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(headers) == 3 and col_idx < 3 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(row_idx, col_idx), text, align=align)

    return table


def add_page_break(document):
    document.add_page_break()


doc = Document()
set_page_geometry(doc.sections[0])
build_styles(doc)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Семестрова контрольна робота з курсу STEM")

add_paragraph(doc, "5 клас", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(
    doc,
    "Укладено на основі навчального плану модулів «Я у Всесвіті», «Я так бачу!» і «Під знаком STEM».",
    style="Small Note",
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
)

add_paragraph(doc, "Прізвище, ім'я _________________________________________________")
add_paragraph(doc, "Клас ____________________    Дата ____________________")

add_paragraph(doc, "Варіант 1", style="Heading 1")
add_paragraph(doc, "1. Тестові завдання", style="Heading 2")
add_numbered_lines(
    doc,
    [
        "Який прилад допомагає спостерігати небесні тіла?\nА. Термометр\nБ. Телескоп\nВ. Мікрофон\nГ. Компас",
        "Що належить до сучасних візуальних медіа?\nА. Анімація\nБ. Формула\nВ. Таблиця множення\nГ. Лінійка",
        "Який із наведених прикладів є штучним знаком?\nА. Темні хмари\nБ. Слід лапи на снігу\nВ. Дорожній знак\nГ. Веселка",
    ],
)

add_paragraph(doc, "2. Теоретичне питання", style="Heading 2")
add_paragraph(doc, "Поясни, яку роль у житті людини відіграють зображення, символи та знаки. Напиши 5-7 речень.")
add_answer_lines(doc, 5)

add_paragraph(doc, "3. Практичне завдання 1", style="Heading 2")
add_paragraph(
    doc,
    "Розподіли подані поняття у три групи: «Космос», «Графіка», «Символи». Запиши всі слова в таблицю.",
)
add_paragraph(
    doc,
    "Слова для розподілу: телескоп, герб, фотографія, орбіта, емблема, анімація, ракета, піктограма, планета, малюнок, печатка, колаж.",
)
add_simple_table(
    doc,
    ["Космос", "Графіка", "Символи"],
    [["", "", ""], ["", "", ""], ["", "", ""], ["", "", ""]],
    widths=[2.1, 2.1, 2.1],
)

add_paragraph(doc, "4. Практичне завдання 2", style="Heading 2")
add_paragraph(doc, "Заповни таблицю «Об'єкт - призначення - модуль».")
add_simple_table(
    doc,
    ["Об'єкт", "Для чого використовується", "До якого модуля належить"],
    [
        ["Телескоп", "", ""],
        ["Герб", "", ""],
        ["Анімація", "", ""],
        ["Ракета", "", ""],
        ["Піктограма", "", ""],
        ["Фотографія", "", ""],
        ["Печатка", "", ""],
        ["Орбіта", "", ""],
    ],
    widths=[1.6, 3.6, 1.3],
)

add_page_break(doc)

add_paragraph(doc, "Варіант 2", style="Heading 1")
add_paragraph(doc, "1. Тестові завдання", style="Heading 2")
add_numbered_lines(
    doc,
    [
        "Що є небезпечним фактором у відкритому космосі?\nА. Космічне сміття\nБ. Дощ\nВ. Туман\nГ. Пісок",
        "Що з наведеного найбільше пов'язане з комп'ютерною графікою?\nА. Створення зображень на комп'ютері\nБ. Вирощування рослин\nВ. Вимірювання температури\nГ. Приготування їжі",
        "Який приклад є природним знаком?\nА. Смайлик\nБ. Літера\nВ. Темні хмари\nГ. Печатка",
    ],
)

add_paragraph(doc, "2. Теоретичне питання", style="Heading 2")
add_paragraph(doc, "Поясни письмово, що таке знакова система та як знаки допомагають людям передавати інформацію. Напиши 5-7 речень.")
add_answer_lines(doc, 5)

add_paragraph(doc, "3. Практичне завдання 1", style="Heading 2")
add_paragraph(doc, "Заповни таблицю «Ризик і захист у космосі».")
add_simple_table(
    doc,
    ["Небезпечний фактор", "Чим загрожує", "Спосіб захисту"],
    [["Холод", "", ""], ["Космічне сміття", "", ""], ["Радіація", "", ""]],
    widths=[2.0, 2.4, 2.1],
)

add_paragraph(doc, "4. Практичне завдання 2", style="Heading 2")
add_paragraph(doc, "Склади план створення інформаційного плаката «Космічні професії майбутнього». У плані має бути не менше 5 кроків.")
for idx in range(1, 7):
    add_paragraph(doc, f"{idx}. ______________________________________________________________")

add_page_break(doc)

add_paragraph(doc, "Тестовий варіант-зразок для учнів", style="Heading 1")
add_paragraph(doc, "Цей варіант можна використати як тренувальний перед контрольною роботою.", style="Small Note")

add_paragraph(doc, "1. Тестові завдання", style="Heading 2")
add_numbered_lines(
    doc,
    [
        "Який прилад допомагає спостерігати небесні тіла?\nА. Термометр\nБ. Телескоп\nВ. Мікрофон\nГ. Компас",
        "Що належить до сучасних візуальних медіа?\nА. Анімація\nБ. Таблиця множення\nВ. Лінійка\nГ. Формула",
        "Який із прикладів є штучним знаком?\nА. Веселка\nБ. Темні хмари\nВ. Дорожній знак\nГ. Слід лапи на снігу",
    ],
)

add_paragraph(doc, "2. Теоретичне питання", style="Heading 2")
add_paragraph(doc, "Поясни, як люди використовують знаки й зображення для передавання інформації.")
add_answer_lines(doc, 5)

add_paragraph(doc, "3. Практичне завдання 1", style="Heading 2")
add_paragraph(doc, "Розподіли слова у три колонки: «Космос», «Графіка», «Символи».")
add_paragraph(
    doc,
    "Слова: ракета, малюнок, печатка, планета, фотографія, знак, телескоп, емблема, орбіта, анімація, герб, піктограма.",
)
add_simple_table(
    doc,
    ["Космос", "Графіка", "Символи"],
    [["", "", ""], ["", "", ""], ["", "", ""], ["", "", ""]],
    widths=[2.1, 2.1, 2.1],
)

add_paragraph(doc, "4. Практичне завдання 2", style="Heading 2")
add_paragraph(doc, "Заповни таблицю «Приклад - група - коротке пояснення».")
add_simple_table(
    doc,
    ["Приклад", "До якої групи належить", "Коротке пояснення"],
    [
        ["Телескоп", "", ""],
        ["Герб", "", ""],
        ["Анімація", "", ""],
        ["Планета", "", ""],
        ["Піктограма", "", ""],
        ["Фотографія", "", ""],
    ],
    widths=[1.5, 1.9, 3.1],
)

add_page_break(doc)

add_paragraph(doc, "Відповіді для вчителя", style="Heading 1")
add_paragraph(doc, "Варіант 1", style="Heading 2")
add_paragraph(doc, "Тести: 1 - Б; 2 - А; 3 - В.")
add_paragraph(
    doc,
    "Теоретичне питання, орієнтовна відповідь: Зображення, символи та знаки допомагають людині передавати інформацію. За допомогою зображень можна показати предмет, явище або подію. Символи передають ідею або значення у короткій формі. Знаки допомагають орієнтуватися в просторі, навчанні та спілкуванні. Наприклад, дорожні знаки попереджають, а герб або емблема показують належність. Тому зображення і знаки важливі в повсякденному житті.",
)
add_simple_table(
    doc,
    ["Космос", "Графіка", "Символи"],
    [
        ["телескоп", "фотографія", "герб"],
        ["орбіта", "анімація", "емблема"],
        ["ракета", "малюнок", "піктограма"],
        ["планета", "колаж", "печатка"],
    ],
    widths=[2.1, 2.1, 2.1],
)
add_simple_table(
    doc,
    ["Об'єкт", "Для чого використовується", "До якого модуля належить"],
    [
        ["Телескоп", "Для спостереження за небесними тілами", "Космос"],
        ["Герб", "Для позначення належності, історії, традицій", "Символи"],
        ["Анімація", "Для створення рухомих зображень", "Графіка"],
        ["Ракета", "Для польотів у космос", "Космос"],
        ["Піктограма", "Для швидкого передавання короткої інформації", "Символи"],
        ["Фотографія", "Для фіксації та передавання зображень", "Графіка"],
        ["Печатка", "Для позначення або підтвердження", "Символи"],
        ["Орбіта", "Для позначення шляху руху тіла в космосі", "Космос"],
    ],
    widths=[1.6, 3.6, 1.3],
)

add_paragraph(doc, "Варіант 2", style="Heading 2")
add_paragraph(doc, "Тести: 1 - А; 2 - А; 3 - В.")
add_paragraph(
    doc,
    "Теоретичне питання, орієнтовна відповідь: Знакова система - це сукупність знаків, за допомогою яких люди передають інформацію. Знаки бувають природні та штучні. Природні знаки існують у природі, наприклад темні хмари можуть означати дощ. Штучні знаки створює людина: букви, цифри, дорожні знаки, емблеми. Вони допомагають швидко повідомляти важливу інформацію. Без знакових систем було б складніше навчатися, спілкуватися й орієнтуватися.",
)
add_simple_table(
    doc,
    ["Небезпечний фактор", "Чим загрожує", "Спосіб захисту"],
    [
        ["Холод", "Переохолодженням", "Теплозахисний скафандр"],
        ["Космічне сміття", "Ударами та пошкодженнями", "Міцний захисний шар, екранування"],
        ["Радіація", "Шкодою для здоров'я", "Спеціальні матеріали захисту, укриття"],
    ],
    widths=[2.0, 2.4, 2.1],
)
add_paragraph(
    doc,
    "Практичне завдання 2, зразок плану: 1. Обрати 3-4 космічні професії. 2. Зібрати коротку інформацію про них. 3. Добрати ілюстрації або піктограми. 4. Розмістити інформацію на плакаті. 5. Виділити потрібні навички для кожної професії. 6. Оформити заголовок і висновок.",
)

add_paragraph(doc, "Варіант-зразок", style="Heading 2")
add_paragraph(doc, "Тести: 1 - Б; 2 - А; 3 - В.")
add_paragraph(
    doc,
    "Теоретичне питання, орієнтовна відповідь: Люди використовують знаки й зображення, щоб швидко передавати інформацію. Малюнок або фотографія допомагає побачити предмет чи явище. Знаки й символи можуть попереджати, пояснювати або позначати щось важливе. Наприклад, дорожній знак допомагає орієнтуватися на дорозі. Герб або емблема показують належність до міста, школи чи команди. Тому знаки та зображення є важливою частиною нашого життя.",
)
add_simple_table(
    doc,
    ["Космос", "Графіка", "Символи"],
    [
        ["ракета", "малюнок", "печатка"],
        ["планета", "фотографія", "знак"],
        ["телескоп", "анімація", "емблема"],
        ["орбіта", "", "герб"],
        ["", "", "піктограма"],
    ],
    widths=[2.1, 2.1, 2.1],
)
add_simple_table(
    doc,
    ["Приклад", "До якої групи належить", "Коротке пояснення"],
    [
        ["Телескоп", "Космос", "Прилад для спостереження за небесними тілами"],
        ["Герб", "Символи", "Передає належність і має символічне значення"],
        ["Анімація", "Графіка", "Створює рухоме зображення"],
        ["Планета", "Космос", "Небесне тіло Сонячної системи або іншої системи"],
        ["Піктограма", "Символи", "Умовний знак для швидкого повідомлення"],
        ["Фотографія", "Графіка", "Фіксує й передає зображення"],
    ],
    widths=[1.5, 1.9, 3.1],
)

doc.sections[-1].start_type = WD_SECTION_START.NEW_PAGE
doc.save(OUTPUT)
print(OUTPUT)
