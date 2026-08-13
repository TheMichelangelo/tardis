from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mikepastula/MyProjects/stem_laboratory")
OUTPUT = ROOT / "kontrolna-stem-2-varianty-bez-vidpovidey.docx"


def set_page_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title.paragraph_format.space_after = Pt(10)

    for name, size, before, after in [
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 14, 14, 6),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_paragraph(document, text="", style="Normal", alignment=None):
    p = document.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    p.add_run(text)
    return p


def add_numbered_question(document, number, text):
    p = document.add_paragraph(style="Normal")
    p.add_run(f"{number}. ").bold = True
    parts = text.split("\n")
    p.add_run(parts[0])
    for extra in parts[1:]:
        p.add_run("\n" + extra)


def add_answer_lines(document, count=5):
    for _ in range(count):
        add_paragraph(document, "_" * 72)


def add_blank_numbered_lines(document, count):
    for idx in range(1, count + 1):
        add_paragraph(document, f"{idx}. " + "_" * 62)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "F1F3F4")

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(row_idx, col_idx), text, align=align)


doc = Document()
set_page_geometry(doc.sections[0])
build_styles(doc)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Семестрова контрольна робота з курсу STEM")

add_paragraph(doc, "5 клас", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(
    doc,
    "На основі навчального плану модулів «Я у Всесвіті», «Я так бачу!», «Під знаком STEM»",
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
)

add_paragraph(doc, "Прізвище, ім'я _________________________________________________")
add_paragraph(doc, "Клас ____________________    Дата ____________________")

add_paragraph(doc, "Варіант 1", style="Heading 1")
add_paragraph(doc, "1. Тестові завдання", style="Heading 2")

variant1_tests = [
    "Який прилад використовують для спостереження за небесними тілами?\nА. Телескоп\nБ. Мікроскоп\nВ. Фотоальбом\nГ. Лупа",
    "Що таке орбіта?\nА. Космічний костюм\nБ. Шлях руху небесного тіла\nВ. Малюнок планети\nГ. Назва ракети",
    "Який об'єкт належить до космічної техніки?\nА. Герб\nБ. Піктограма\nВ. Ракета\nГ. Печатка",
    "Що допомагає створювати рухомі зображення?\nА. Анімація\nБ. Орбіта\nВ. Телескоп\nГ. Скафандр",
    "Який із наведених прикладів є штучним знаком?\nА. Веселка\nБ. Темні хмари\nВ. Дорожній знак\nГ. Слід лапи",
    "Для чого космонавту потрібен скафандр?\nА. Для прикрашання\nБ. Для захисту в космосі\nВ. Для малювання\nГ. Для подачі сигналів",
]

for idx, question in enumerate(variant1_tests, start=1):
    add_numbered_question(doc, idx, question)

add_paragraph(doc, "2. Теоретичне питання", style="Heading 2")
add_paragraph(doc, "Поясни, яку роль відіграють кольори у графіці, символах і повсякденному житті людини.")
add_answer_lines(doc, 5)

add_paragraph(doc, "3. Практичне завдання 1", style="Heading 2")
add_paragraph(doc, "Заповни таблицю про космос.")
add_table(
    doc,
    ["Об'єкт / поняття", "Що це або для чого використовується", "До якої теми належить"],
    [
        ["Телескоп", "", ""],
        ["Орбіта", "", ""],
        ["Ракета", "", ""],
        ["Скафандр", "", ""],
        ["Планета", "", ""],
        ["Космічна станція", "", ""],
    ],
    [1.7, 3.6, 1.2],
)

add_paragraph(doc, "4. Практичне завдання 2", style="Heading 2")
add_paragraph(doc, "Склади план розробки моделі ракети. Запиши 5-7 послідовних кроків.")
add_blank_numbered_lines(doc, 7)

add_paragraph(doc, "5. Практичне завдання 3", style="Heading 2")
add_answer_lines(doc, 5)

doc.add_page_break()

add_paragraph(doc, "Прізвище, ім'я _________________________________________________")
add_paragraph(doc, "Клас ____________________    Дата ____________________")

add_paragraph(doc, "Варіант 2", style="Heading 1")
add_paragraph(doc, "1. Тестові завдання", style="Heading 2")

variant2_tests = [
    "Який приклад є природним знаком?\nА. Смайлик\nБ. Буква\nВ. Темні хмари\nГ. Печатка",
    "Що належить до графічних знакових систем?\nА. Герб\nБ. Планета\nВ. Орбіта\nГ. Скафандр",
    "Що таке піктограма?\nА. Небесне тіло\nБ. Умовний малюнок-знак\nВ. Космічний апарат\nГ. Вид фарби",
    "Яке з наведеного пов'язане з фотографією?\nА. Фотоапарат\nБ. Телескоп\nВ. Герб\nГ. Печатка",
    "Що допомагає передавати інформацію в символічній формі?\nА. Знаки і символи\nБ. Лише планети\nВ. Лише кольори\nГ. Тільки ракети",
    "До якого модуля належить тема фотографії та анімації?\nА. Людина - природа\nБ. Людина - образ\nВ. Людина - знак\nГ. Людина - техніка",
]

for idx, question in enumerate(variant2_tests, start=1):
    add_numbered_question(doc, idx, question)

add_paragraph(doc, "2. Теоретичне питання", style="Heading 2")
add_paragraph(doc, "Поясни, як кольори допомагають людині передавати настрій, значення та інформацію в малюнках, символах і фотографіях.")
add_answer_lines(doc, 5)

add_paragraph(doc, "3. Практичне завдання 1", style="Heading 2")
add_paragraph(doc, "Заповни таблицю про знаки і символи.")
add_table(
    doc,
    ["Знак / символ", "Що означає або для чого використовується", "До якої теми належить"],
    [
        ["Герб", "", ""],
        ["Печатка", "", ""],
        ["Піктограма", "", ""],
        ["Дорожній знак", "", ""],
        ["Буква", "", ""],
        ["Емблема", "", ""],
    ],
    [1.7, 3.6, 1.2],
)

add_paragraph(doc, "4. Практичне завдання 2", style="Heading 2")
add_paragraph(doc, "Склади план розробки моделі фотоапарата. Запиши 5-7 послідовних кроків.")
add_blank_numbered_lines(doc, 7)

add_paragraph(doc, "5. Практичне завдання 3", style="Heading 2")
add_answer_lines(doc, 5)

doc.save(OUTPUT)
print(OUTPUT)
