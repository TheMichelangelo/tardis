from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/mikepastula/MyProjects/stem_laboratory")
OUTPUT = ROOT / "znakovi-systemy-ta-profesiyi.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section, top=2.0, bottom=2.0, left=2.2, right=2.2):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


document = Document()

section = document.sections[0]
set_page_margins(section)

styles = document.styles

normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)

title_style = styles["Title"]
title_style.font.name = "Arial"
title_style.font.size = Pt(18)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

heading1 = styles["Heading 1"]
heading1.font.name = "Arial"
heading1.font.size = Pt(14)
heading1.font.bold = True
heading1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

if "Task Body" not in styles:
    task_body = styles.add_style("Task Body", WD_STYLE_TYPE.PARAGRAPH)
else:
    task_body = styles["Task Body"]
task_body.base_style = normal
task_body.font.name = "Arial"
task_body.font.size = Pt(11)
task_body.paragraph_format.space_after = Pt(6)
task_body.paragraph_format.line_spacing = 1.15

if "Task Bullet" not in styles:
    task_bullet = styles.add_style("Task Bullet", WD_STYLE_TYPE.PARAGRAPH)
else:
    task_bullet = styles["Task Bullet"]
task_bullet.base_style = normal
task_bullet.font.name = "Arial"
task_bullet.font.size = Pt(11)
task_bullet.paragraph_format.left_indent = Cm(0.6)
task_bullet.paragraph_format.space_after = Pt(3)
task_bullet.paragraph_format.line_spacing = 1.1


title = document.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Урок: Знакові системи та професії, пов'язані з ними")

subtitle = document.add_paragraph(style="Task Body")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Навчальний матеріал із 7 завданнями для роботи в класі та вдома")
subtitle_run.italic = True


def add_heading(text):
    p = document.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    p.add_run(text)


def add_body(text, bold_prefix=None):
    p = document.add_paragraph(style="Task Body")
    if bold_prefix and text.startswith(bold_prefix):
        prefix, rest = text[: len(bold_prefix)], text[len(bold_prefix) :]
        run = p.add_run(prefix)
        run.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)


def add_numbered_line(number, text):
    p = document.add_paragraph(style="Task Body")
    num = p.add_run(f"{number}. ")
    num.bold = True
    p.add_run(text)


def add_bullet(text):
    p = document.add_paragraph(style="Task Bullet")
    p.add_run("• ")
    p.add_run(text)


add_heading("Завдання 1. Прочитай визначення професій")

professions = [
    (
        "Програміст",
        "це фахівець, який створює комп'ютерні програми, сайти, ігри та застосунки за допомогою спеціальних мов програмування.",
        "продумує, як має працювати програма, записує команди для комп'ютера, перевіряє помилки та вдосконалює результат.",
    ),
    (
        "Перекладач",
        "це фахівець, який перекладає усне або письмове мовлення з однієї мови на іншу.",
        "уважно читає або слухає текст, розуміє його зміст, добирає точні слова іншою мовою та допомагає людям порозумітися.",
    ),
    (
        "Банківський касир",
        "це працівник банку, який приймає, видає та рахує гроші, а також оформлює прості фінансові операції.",
        "перевіряє суми, заповнює таблиці та документи, уважно працює з числами, приймає платежі або видає кошти клієнтам.",
    ),
    (
        "Вчитель математики",
        "це педагог, який навчає учнів розв'язувати задачі, виконувати обчислення та логічно мислити.",
        "пояснює нові теми, добирає приклади і задачі, перевіряє відповіді учнів та допомагає зрозуміти правила і формули.",
    ),
]

for index, (name, definition, work) in enumerate(professions, start=1):
    p = document.add_paragraph(style="Task Body")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{index}. {name}")
    run.bold = True
    p.add_run(f" — {definition}")
    add_body(f"Що робить {name.lower()}: {work}", bold_prefix=f"Що робить {name.lower()}: ")


add_heading("Завдання 2. Спробуй себе в ролі програміста")
add_body('Намалюй просту блок-схему "Як зібратися до школи зранку".')
add_bullet("У блок-схемі має бути не менше 6 кроків.")
add_bullet("Використай слова: початок, дія, перевірка, кінець.")
add_bullet("Приклад кроків: прокинувся, вмився, поснідав, перевірив рюкзак, одягнувся, вийшов до школи.")
add_body("Час виконання: 5-7 хвилин.", bold_prefix="Час виконання: ")


add_heading("Завдання 3. Спробуй себе в ролі перекладача")
add_body("Переклади українською мовою 3 речення:")
add_numbered_line(1, "I like math and science.")
add_numbered_line(2, "My friend reads an interesting book.")
add_numbered_line(3, "We learn new symbols at school.")
add_body("Після перекладу підкресли слова, значення яких ти знав або знайшов самостійно.")
add_body("Час виконання: 5-7 хвилин.", bold_prefix="Час виконання: ")


add_heading("Завдання 4. Спробуй себе в ролі банківського касира")
add_body("Заповни таблицю за зразком.")
add_body("Дані:", bold_prefix="Дані:")
add_bullet("У понеділок до каси надійшло 500 грн, видали 200 грн.")
add_bullet("У вівторок надійшло 300 грн, видали 150 грн.")
add_bullet("У середу надійшло 450 грн, видали 250 грн.")

table = document.add_table(rows=4, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

headers = ["День", "Надійшло (грн)", "Видано (грн)", "Залишок (грн)"]
for idx, text in enumerate(headers):
    cell = table.cell(0, idx)
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "D9EAF7")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

rows = [
    ["Понеділок", "500", "200", "______"],
    ["Вівторок", "300", "150", "______"],
    ["Середа", "450", "250", "______"],
]

for row_idx, row in enumerate(rows, start=1):
    for col_idx, text in enumerate(row):
        cell = table.cell(row_idx, col_idx)
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body("Обчисли залишок грошей за кожен день.")
add_body("Час виконання: 5-7 хвилин.", bold_prefix="Час виконання: ")


add_heading("Завдання 5. Спробуй себе в ролі вчителя математики")
add_body("Склади і запиши:")
add_numbered_line(1, "Один приклад на додавання.")
add_numbered_line(2, "Один приклад на віднімання.")
add_numbered_line(3, "Одну коротку задачу для однокласника.")
add_body("Після цього обміняйся завданням із сусідом по парті та перевір його відповідь.")
add_body("Час виконання: 5-7 хвилин.", bold_prefix="Час виконання: ")


add_heading("Завдання 6. Переваги та недоліки професій")
add_body("Запиши для кожної професії по 2 переваги і по 2 недоліки:")
add_bullet("програміст;")
add_bullet("перекладач;")
add_bullet("банківський касир;")
add_bullet("вчитель математики.")
add_body("Спробуй подумати, що в цій професії цікаво, а що може бути складним.")


add_heading("Завдання 7. Домашнє завдання")
add_body('Напиши твір на 10 речень на тему: "Чого потрібно навчатись сьогодні, щоб бути успішним завтра".')
add_body("У творі можна написати:")
add_bullet("які знання та вміння важливі для сучасної людини;")
add_bullet("чому потрібно вчитися працювати з інформацією, числами, мовами та технологіями;")
add_bullet("які професії можуть бути потрібні в майбутньому.")


document.add_section(WD_SECTION.CONTINUOUS)
document.save(OUTPUT)
print(OUTPUT)
